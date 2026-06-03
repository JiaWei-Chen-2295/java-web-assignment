from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


INPUT = Path("实验报告_Session与过滤器.docx")
OUTPUT = Path("实验报告_Session与过滤器_论文格式版.docx")


def set_font(run, east_asia="宋体", ascii_font="Times New Roman", size=None, bold=None, color=RGBColor(0, 0, 0)):
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color
    run.font.name = ascii_font
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), east_asia)
    rfonts.set(qn("w:ascii"), ascii_font)
    rfonts.set(qn("w:hAnsi"), ascii_font)


def set_style_font(style, east_asia, ascii_font, size, bold=None):
    font = style.font
    font.name = ascii_font
    font.size = Pt(size)
    if bold is not None:
        font.bold = bold
    rpr = style._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), east_asia)
    rfonts.set(qn("w:ascii"), ascii_font)
    rfonts.set(qn("w:hAnsi"), ascii_font)


def style_paragraph(paragraph, *, align=None, first_indent=True, before=0, after=6, line=1.5):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    fmt.line_spacing = line
    fmt.first_line_indent = Pt(24) if first_indent else Pt(0)
    if align is not None:
        paragraph.alignment = align


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr, fld_sep, text, fld_end])
    set_font(run, size=10.5)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def is_code_paragraph(paragraph):
    if not paragraph.text.strip():
        return False
    return any((run.font.name or "").lower() in {"consolas", "courier new"} for run in paragraph.runs)


def main():
    doc = Document(INPUT)

    section = doc.sections[0]
    section.start_type = WD_SECTION.NEW_PAGE
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.6)
    section.header_distance = Cm(1.5)
    section.footer_distance = Cm(1.75)

    # Clear old footer content and add centered page number.
    for p in section.footer.paragraphs:
        p._element.clear_content()
    add_page_number(section.footer.paragraphs[0])

    styles = doc.styles
    set_style_font(styles["Normal"], "宋体", "Times New Roman", 12)
    styles["Normal"].paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    styles["Normal"].paragraph_format.line_spacing = 1.5
    styles["Normal"].paragraph_format.space_after = Pt(6)

    set_style_font(styles["Heading 1"], "黑体", "Times New Roman", 15, True)
    styles["Heading 1"].font.color.rgb = RGBColor(0, 0, 0)
    styles["Heading 1"].paragraph_format.space_before = Pt(12)
    styles["Heading 1"].paragraph_format.space_after = Pt(6)
    styles["Heading 1"].paragraph_format.first_line_indent = Pt(0)

    set_style_font(styles["Heading 2"], "黑体", "Times New Roman", 14, True)
    styles["Heading 2"].font.color.rgb = RGBColor(0, 0, 0)
    styles["Heading 2"].paragraph_format.space_before = Pt(8)
    styles["Heading 2"].paragraph_format.space_after = Pt(4)
    styles["Heading 2"].paragraph_format.first_line_indent = Pt(0)

    set_style_font(styles["Heading 3"], "黑体", "Times New Roman", 12, True)
    styles["Heading 3"].font.color.rgb = RGBColor(0, 0, 0)
    styles["Heading 3"].paragraph_format.space_before = Pt(6)
    styles["Heading 3"].paragraph_format.space_after = Pt(3)
    styles["Heading 3"].paragraph_format.first_line_indent = Pt(0)

    for idx, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip()
        style_name = paragraph.style.name
        if not text:
            continue

        if idx == 0:
            style_paragraph(paragraph, align=WD_ALIGN_PARAGRAPH.CENTER, first_indent=False, before=0, after=10, line=1.2)
            for run in paragraph.runs:
                set_font(run, "黑体", "Times New Roman", 22, True)
        elif idx == 1:
            style_paragraph(paragraph, align=WD_ALIGN_PARAGRAPH.CENTER, first_indent=False, before=0, after=18, line=1.2)
            for run in paragraph.runs:
                set_font(run, "宋体", "Times New Roman", 15, False)
        elif style_name == "Heading 1":
            style_paragraph(paragraph, first_indent=False, before=12, after=6, line=1.5)
            for run in paragraph.runs:
                set_font(run, "黑体", "Times New Roman", 15, True)
        elif style_name == "Heading 2":
            style_paragraph(paragraph, first_indent=False, before=8, after=4, line=1.5)
            for run in paragraph.runs:
                set_font(run, "黑体", "Times New Roman", 14, True)
        elif style_name.startswith("List"):
            style_paragraph(paragraph, first_indent=False, before=0, after=3, line=1.5)
            paragraph.paragraph_format.left_indent = Pt(24)
            for run in paragraph.runs:
                set_font(run, "宋体", "Times New Roman", 12)
        elif is_code_paragraph(paragraph):
            style_paragraph(paragraph, first_indent=False, before=3, after=6, line=1.0)
            paragraph.paragraph_format.left_indent = Pt(12)
            for run in paragraph.runs:
                set_font(run, "宋体", "Consolas", 9.5)
        elif text.startswith("【图") or text.startswith("图"):
            style_paragraph(paragraph, align=WD_ALIGN_PARAGRAPH.CENTER, first_indent=False, before=3, after=8, line=1.2)
            for run in paragraph.runs:
                set_font(run, "宋体", "Times New Roman", 10.5)
        else:
            style_paragraph(paragraph, first_indent=True, before=0, after=6, line=1.5)
            for run in paragraph.runs:
                set_font(run, "宋体", "Times New Roman", 12)

    for table in doc.tables:
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        for r_idx, row in enumerate(table.rows):
            for cell in row.cells:
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                if r_idx == 0:
                    shade_cell(cell, "D9EAF7")
                for paragraph in cell.paragraphs:
                    style_paragraph(paragraph, first_indent=False, before=0, after=0, line=1.2)
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if len(paragraph.text) < 24 else WD_ALIGN_PARAGRAPH.LEFT
                    for run in paragraph.runs:
                        set_font(run, "宋体", "Times New Roman", 10.5, r_idx == 0 or None)

    core = doc.core_properties
    core.title = "Java Web 实验报告：Session 会话管理与 Servlet 过滤器应用"
    core.subject = "中文论文格式排版"

    doc.save(OUTPUT)
    print(OUTPUT.resolve())


if __name__ == "__main__":
    main()
