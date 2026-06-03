# -*- coding: utf-8 -*-
"""Generate Session & Filter lab report as DOCX."""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from pathlib import Path

OUTPUT = Path(__file__).resolve().parents[1] / "实验报告_Session与过滤器.docx"


def set_doc_font(doc):
    style = doc.styles["Normal"]
    style.font.name = "宋体"
    style.font.size = Pt(12)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")


def add_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(22)
    run.font.name = "黑体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")


def add_h1(doc, text):
    doc.add_heading(text, level=1)


def add_h2(doc, text):
    doc.add_heading(text, level=2)


def add_h3(doc, text):
    doc.add_heading(text, level=3)


def add_para(doc, text):
    doc.add_paragraph(text)


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_code(doc, code):
    p = doc.add_paragraph()
    run = p.add_run(code)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    pf = p.paragraph_format
    pf.left_indent = Cm(0.5)
    pf.space_before = Pt(6)
    pf.space_after = Pt(6)


def add_screenshot_placeholder(doc, caption):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"【{caption}】")
    run.italic = True
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    doc.add_paragraph()


def build():
    doc = Document()
    set_doc_font(doc)

    # margins
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)

    add_title(doc, "Java Web 实验报告")
    add_para(doc, "Session 会话管理与 Servlet 过滤器应用")
    add_para(doc, "")
    add_para(doc, "实验主题：Session 访问统计、Session 购物车、编码过滤器、权限过滤器、注册验证过滤器、监听器在线统计")

    # 1
    add_h1(doc, "1 实验目的")
    add_para(doc, "通过本次实验掌握 HttpSession 的创建与属性存取、基于 Session 的购物车模拟、Servlet Filter 过滤器链的拦截与放行机制，以及 HttpSessionListener 监听会话生命周期实现在线用户统计的方法。")

    # 2
    add_h1(doc, "2 实验环境")
    table = doc.add_table(rows=6, cols=2)
    table.style = "Table Grid"
    rows = [
        ("操作系统", "Windows 11"),
        ("JDK", "17"),
        ("Web 容器", "Apache Tomcat 10+（Jakarta EE）"),
        ("构建工具", "Maven 3.x"),
        ("技术规范", "Jakarta Servlet 6.x + JSP 3.x"),
    ]
    for i, (k, v) in enumerate(rows):
        table.rows[i].cells[0].text = k
        table.rows[i].cells[1].text = v
    doc.add_paragraph()

    # 3 Session access count
    add_h1(doc, "3 Session 访问次数统计")
    add_h2(doc, "3.1 设计思路")
    add_para(doc, "每次用户访问页面时，Servlet 通过 request.getSession() 获取（或创建）当前浏览器对应的 Session 对象，从 Session 中读取名为 accesscount 的整数属性：若不存在则置为 1，否则自增 1 后写回。同一 Session 内反复刷新页面，计数持续累加；关闭浏览器或 Session 过期后计数重置。")
    add_h2(doc, "3.2 核心代码")
    add_code(doc, """@WebServlet("/accessCount")
protected void doGet(...) {
    HttpSession session = request.getSession();
    Integer count = (Integer) session.getAttribute("accesscount");
    count = (count == null) ? 1 : count + 1;
    session.setAttribute("accesscount", count);
    // 输出当前访问次数与 Session ID
}""")
    add_h2(doc, "3.3 文件结构")
    add_code(doc, """ch12-session-counter/
├── pom.xml
└── src/main/java/.../AccessCountServlet.java""")
    add_h2(doc, "3.4 运行截图")
    add_screenshot_placeholder(doc, "图3-1 多次刷新后访问次数递增的页面效果")

    # 4 Session cart
    add_h1(doc, "4 Session 模拟购物车")
    add_h2(doc, "4.1 设计思路")
    add_para(doc, "购物车数据以 Map<String, CartItem> 形式存入 Session，键为商品 ID。用户点击「加入购物车」时，Servlet 从 Session 取出 Map，若商品已存在则数量 +1，否则新建 CartItem。模块提供三种提交策略对比：")
    add_bullets(doc, [
        "直接 Forward：刷新会重复提交 POST，数量异常增加；",
        "Redirect（PRG 模式）：POST 后重定向 GET，刷新不会重复加购；",
        "Session 防重：用 requestId 集合记录已处理请求，拦截重复 POST。",
    ])
    add_h2(doc, "4.2 核心代码")
    add_code(doc, """// CartSupport：Session 购物车读写
public static void addProduct(HttpSession session, String mode, Product p) {
    Map<String, CartItem> cart = getCart(session, mode);
    CartItem item = cart.get(p.getId());
    if (item == null) cart.put(p.getId(), new CartItem(p));
    else item.increment();
}

// RedirectCartServlet：PRG 防重复提交
CartSupport.addProduct(session, "redirect", product);
response.sendRedirect(contextPath + "/shop?tab=redirect");""")
    add_h2(doc, "4.3 文件结构")
    add_code(doc, """ch12-session-cart/
├── model/   CartItem.java, Product.java
├── service/ CartSupport.java, CatalogService.java
├── *CartServlet.java, ShopServlet.java
└── webapp/WEB-INF/shop.jsp""")
    add_h2(doc, "4.4 运行截图")
    add_screenshot_placeholder(doc, "图4-1 商城页面三种购物车模式对比")
    add_screenshot_placeholder(doc, "图4-2 加入商品后 Session 中购物车列表")

    # 5 Encoding filter
    add_h1(doc, "5 过滤器解决中文乱码")
    add_h2(doc, "5.1 设计思路")
    add_para(doc, "Tomcat 默认以 ISO-8859-1 解析 POST 请求体，直接读取中文参数会出现乱码。EncodingFilter 在请求到达 Servlet 之前统一设置 request/response 字符编码为 UTF-8，并指定 HTML 响应的 Content-Type，使注册表单提交的中文昵称能正确显示。")
    add_h2(doc, "5.2 核心代码")
    add_code(doc, """public void doFilter(ServletRequest req, ServletResponse resp, FilterChain chain) {
    req.setCharacterEncoding("UTF-8");
    resp.setCharacterEncoding("UTF-8");
    ((HttpServletResponse) resp).setContentType("text/html;charset=UTF-8");
    chain.doFilter(req, resp);  // 放行到 RegServler
}""")
    add_para(doc, "web.xml 中将 EncodingFilter 映射到 /reg 路径，register.html 表单提交中文昵称即可验证效果。")
    add_h2(doc, "5.3 文件结构")
    add_code(doc, """ch14-filter-intro/
├── EncodingFilter.java
├── RegServler.java
├── webapp/register.html
└── WEB-INF/web.xml  （filter-mapping → /reg）""")
    add_h2(doc, "5.4 运行截图")
    add_screenshot_placeholder(doc, "图5-1 提交中文昵称后注册结果页正确显示")

    # 6 Permission filter
    add_h1(doc, "6 过滤器进行用户权限访问控制")
    add_h2(doc, "6.1 设计思路")
    add_para(doc, "采用「登录校验 + 角色鉴权」两级过滤器链：")
    add_bullets(doc, [
        "LoginValidatorFilter：拦截 POST /login，比对 Session 中的验证码，错误则 forward 回登录页；",
        "LoginServlet：验证账号密码，将 username、role 写入 Session；",
        "PermissionFilter：拦截 /admin/* 与 /user/*，未登录重定向登录页，非 admin 角色访问管理页返回 403。",
    ])
    add_h2(doc, "6.2 核心代码")
    add_code(doc, """// PermissionFilter
HttpSession session = httpReq.getSession(false);
if (session == null || session.getAttribute("username") == null) {
    httpResp.sendRedirect(contextPath + "/login.jsp?error=2");
    return;
}
if (uri.contains("/admin/") && !"admin".equals(session.getAttribute("role"))) {
    httpResp.sendError(403, "权限不足");
    return;
}
chain.doFilter(request, response);""")
    add_h2(doc, "6.3 文件结构")
    add_code(doc, """ch14-filter-intro/
├── demo/filter/LoginValidatorFilter.java
├── demo/filter/PermissionFilter.java
├── controller/LoginServlet.java
├── webapp/login.jsp, admin/admin.jsp, user/user.jsp
└── WEB-INF/web.xml""")
    add_h2(doc, "6.4 运行截图")
    add_screenshot_placeholder(doc, "图6-1 管理员登录后进入 admin 页面")
    add_screenshot_placeholder(doc, "图6-2 普通用户访问 admin 返回 403")

    # 7 Registration validation filter
    add_h1(doc, "7 过滤器验证用户注册")
    add_h2(doc, "7.1 设计思路")
    add_para(doc, "ResValidatorFilter 在 RegisterServlet 处理之前对表单参数做服务端校验：用户名非空且长度 3~20 位、密码非空且长度 6~20 位。校验失败时将错误信息与已填字段 forward 回 register.jsp，避免无效数据进入业务层。")
    add_h2(doc, "7.2 核心代码")
    add_code(doc, """public void doFilter(...) {
    String username = request.getParameter("username");
    String password = request.getParameter("password");
    if (isBlank(username)) errorMsg = "用户名不能为空！";
    else if (username.length() < 3 || username.length() > 20)
        errorMsg = "用户名长度必须在 3~20 位之间！";
    else if (password.length() < 6) errorMsg = "密码长度不足！";

    if (errorMsg != null) {
        request.setAttribute("errorMsg", errorMsg);
        request.getRequestDispatcher("/register.jsp").forward(...);
        return;
    }
    chain.doFilter(request, response);
}""")
    add_h2(doc, "7.3 文件结构")
    add_code(doc, """ch14-filter-intro/
├── ResValidatorFilter.java
├── RegisterServlet.java
├── webapp/register.jsp
└── WEB-INF/web.xml  （filter-mapping → /register）""")
    add_h2(doc, "7.4 运行截图")
    add_screenshot_placeholder(doc, "图7-1 用户名过短时的过滤器拦截提示")
    add_screenshot_placeholder(doc, "图7-2 校验通过后 RegisterServlet 注册成功页")

    # 8 Listener online count
    add_h1(doc, "8 监听器完成在线人数统计")
    add_h2(doc, "8.1 设计思路")
    add_para(doc, "ch11-login-keeper 模块通过 HttpSessionListener 与 SessionManager 协作实现在线会话统计：")
    add_bullets(doc, [
        "用户登录成功后，LoginServlet 调用 SessionManager.login(sessionId, username) 将 Session 登记到内存 Map；",
        "LoginObserver 监听 sessionDestroyed 事件，Session 过期或用户退出时自动从 Map 移除，防止脏数据；",
        "welcome.jsp 读取 activeSessions.size() 展示当前活跃会话数，即在线用户数。",
    ])
    add_h2(doc, "8.2 核心代码")
    add_code(doc, """@WebListener
public class LoginObserver implements HttpSessionListener {
    public void sessionDestroyed(HttpSessionEvent se) {
        SessionManager.getInstance().logout(se.getSession().getId());
    }
}

// LoginServlet 登录成功时
sessionManager.login(session.getId(), username);

// welcome.jsp 展示
activeSessions.size()  // 活跃会话 / 在线人数""")
    add_h2(doc, "8.3 文件结构")
    add_code(doc, """ch11-login-keeper/
├── LoginObserver.java      （HttpSessionListener）
├── SessionManager.java     （ConcurrentHashMap 维护在线 Session）
├── LoginServlet.java
└── webapp/welcome.jsp      （展示活跃会话统计）""")
    add_h2(doc, "8.4 运行截图")
    add_screenshot_placeholder(doc, "图8-1 welcome 控制面板显示活跃会话数量")
    add_screenshot_placeholder(doc, "图8-2 多用户登录后会话列表")

    # 9 Summary
    add_h1(doc, "9 实验总结")
    add_para(doc, "本次实验围绕 Session 与 Filter 两条主线展开：Session 负责在服务器端为每个浏览器维持独立状态，可用于访问计数、购物车等场景；Filter 作为请求链路上的横切组件，可集中处理编码、鉴权、表单校验等通用逻辑，避免在每个 Servlet 中重复编写。监听器则补充了 Session 生命周期的自动化管理，配合 SessionManager 实现可靠的在线用户统计。")
    add_para(doc, "实验中需注意：Filter 的 mapping 顺序决定拦截范围；POST 请求必须在读取参数前设置 UTF-8 编码；PRG 模式可有效避免购物车重复提交；权限控制应同时校验「是否登录」与「角色是否匹配」两层条件。")

    doc.save(OUTPUT)
    print(f"Generated: {OUTPUT}")


if __name__ == "__main__":
    build()
